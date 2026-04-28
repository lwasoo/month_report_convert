from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
import urllib.error
import urllib.request

from docx import Document

from report_converter.common import log, normalize_text


COMPANY_SUFFIXES = [
    "股份有限公司",
    "有限责任公司",
    "集团有限公司",
    "科技有限公司",
    "技术有限公司",
    "电子有限公司",
    "实业有限公司",
    "贸易有限公司",
    "国际有限公司",
    "分公司",
    "有限公司",
    "律师事务所",
    "律所",
]
COMPANY_SUFFIX_PATTERN = "|".join(sorted((re.escape(x) for x in COMPANY_SUFFIXES), key=len, reverse=True))
GENERIC_COMPANY_VALUES = {
    "本公司",
    "该公司",
    "我司",
    "贵司",
    "公司",
    "集团",
    "集团公司",
    "某公司",
}
SENTENCE_LIKE_TOKENS = {"采取", "进行", "完成", "存在", "涉及", "相关", "推动", "处理", "开展", "公司应", "公司就"}
COMPANY_STOPWORDS = {
    "其中",
    "协助",
    "新增",
    "需要",
    "并约",
    "申请人",
    "被申请人",
    "妥善",
    "补充",
    "提交",
    "办理",
    "联系",
    "梳理",
    "平台",
    "竞争对手",
    "供应商",
    "客户",
    "终止与",
    "再向",
    "接单",
    "口头",
    "协调",
    "利用",
    "认为",
    "体检",
    "组织",
    "通过",
    "资料",
}

PLACEHOLDER_PATTERNS: list[tuple[str, str]] = [
    ("COMPANY", rf"[\u4e00-\u9fffA-Za-z0-9（）()·&\-_]{{2,18}}(?:{COMPANY_SUFFIX_PATTERN})"),
    ("COMPANY", r"\b[A-Z][A-Za-z&.\-]+(?:\s+[A-Z][A-Za-z&.\-]+){0,5}\s+(?:Legal|Technology|Technologies|Group|Holdings|Partners|Botts|Ltd|LTD|Inc|INC|LLC|PTE|Corp|Corporation|Company)\b"),
    ("COMPANY", r"[A-Za-z\u4e00-\u9fff]{2,24}(?:律师事务所|律所)"),
    ("TITLE", r"《[^》]{2,80}》"),
    ("AMOUNT", r"(?:人民币|USD|RMB|美元)?\s*\d[\d,]*(?:\.\d+)?\s*(?:亿元|万元|元|万美元|美元)"),
    ("ACCOUNT", r"\b\d{12,24}\b"),
    ("EMAIL", r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"),
    ("PHONE", r"\b1[3-9]\d{9}\b"),
    ("CASE", r"\b(?:[A-Z]{2,}-)?(?:TECH-)?\d{3,}\b"),
    ("CODE", r"\b[A-Z]{2,}(?:[-_][A-Z0-9]{2,})+\b"),
]

CONTEXT_RULES: list[tuple[str, str]] = [
    ("PARTY", r"(?<=甲方[：:])[^，,。；;\n]{2,60}"),
    ("PARTY", r"(?<=乙方[：:])[^，,。；;\n]{2,60}"),
    ("PARTY", r"(?<=丙方[：:])[^，,。；;\n]{2,60}"),
    ("PROJECT", r"(?<=项目名称为)[^，,。；;\n]{2,80}"),
    ("PROJECT", r"(?<=项目名称[：:])[^，,。；;\n]{2,80}"),
    ("PROJECT", r"(?<=项目[：:])[^，,。；;\n]{2,80}"),
    ("CASE", r"(?<=案号[：:])[^，,。；;\n]{2,60}"),
    ("CASE", r"(?<=合同编号[：:])[^，,。；;\n]{2,60}"),
    ("SUPPLIER", r"(?<=供应商[：:])[^，,。；;\n]{2,80}"),
    ("CUSTOMER", r"(?<=客户[：:])[^，,。；;\n]{2,80}"),
]

PERSON_CONTEXT_PATTERNS: list[tuple[str, str]] = [
    (
        "PERSON",
        r"(?:申请人|被申请人|申请执行人|被申请执行人|原告|被告|仲裁员|代理人|联系人|员工|涉案人员)[：:\s]*"
        r"((?:欧阳|司马|上官|诸葛|皇甫|尉迟|公孙|长孙|慕容|司徒|夏侯|东方|独孤|南宫|闻人|令狐|轩辕|赵|钱|孙|李|周|吴|郑|王|冯|陈|褚|卫|蒋|沈|韩|杨|朱|秦|尤|许|何|吕|施|张|孔|曹|严|华|金|魏|陶|姜|戚|谢|邹|喻|柏|窦|章|云|苏|潘|葛|范|彭|郎|鲁|韦|昌|马|苗|凤|花|方|俞|任|袁|柳|鲍|史|唐|费|廉|岑|薛|雷|贺|倪|汤|滕|殷|罗|毕|郝|邬|安|常|乐|于|时|傅|皮|卞|齐|康|伍|余|元|卜|顾|孟|平|黄|和|穆|萧|尹)[一-龥某]{1,2})"
    ),
]


@dataclass
class ReplacementItem:
    placeholder: str
    original: str
    category: str
    enabled: bool = True
    source: str = "auto"


def sanitize_docx(
    input_path: Path,
    output_path: Path,
    mapping_path: Path,
    custom_terms: list[str] | None = None,
    use_llm_assist: bool = False,
    model: str = "qwen2.5:7b-instruct-q4_K_M",
    ollama_url: str = "http://127.0.0.1:11434",
    timeout_sec: int = 120,
    retries: int = 2,
) -> dict[str, Any]:
    if input_path.suffix.lower() != ".docx":
        raise ValueError("当前仅支持 .docx 文件。")
    payload = scan_docx(
        input_path=input_path,
        custom_terms=custom_terms or [],
        use_llm_assist=use_llm_assist,
        model=model,
        ollama_url=ollama_url,
        timeout_sec=timeout_sec,
        retries=retries,
        existing_mapping_path=mapping_path if mapping_path.exists() else None,
    )
    apply_mapping_to_docx(input_path, output_path, payload, mapping_path)
    log(f"脱敏完成: {output_path}")
    log(f"映射文件已输出: {mapping_path}")
    return payload


def scan_docx(
    input_path: Path,
    custom_terms: list[str] | None = None,
    use_llm_assist: bool = False,
    model: str = "qwen2.5:7b-instruct-q4_K_M",
    ollama_url: str = "http://127.0.0.1:11434",
    timeout_sec: int = 120,
    retries: int = 2,
    existing_mapping_path: Path | None = None,
    existing_payload: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if input_path.suffix.lower() != ".docx":
        raise ValueError("当前仅支持 .docx 文件。")
    doc = Document(str(input_path))
    output_path = input_path.with_name(f"{input_path.stem}_脱敏.docx")
    mapping_path = existing_mapping_path or input_path.with_name(f"{input_path.stem}_映射.json")
    payload = build_mapping_payload(
        doc,
        input_path=input_path,
        output_path=output_path,
        mapping_path=mapping_path,
        custom_terms=custom_terms or [],
        use_llm_assist=use_llm_assist,
        model=model,
        ollama_url=ollama_url,
        timeout_sec=timeout_sec,
        retries=retries,
        existing_entries_override=existing_payload.get("entries", []) if existing_payload else None,
    )
    if existing_payload:
        payload["source_file"] = str(input_path)
    return payload


def apply_mapping_to_docx(
    input_path: Path,
    output_path: Path,
    payload: dict[str, Any],
    mapping_path: Path | None = None,
) -> None:
    if input_path.suffix.lower() != ".docx":
        raise ValueError("当前仅支持 .docx 文件。")
    doc = Document(str(input_path))
    apply_replacements(doc, mapping_entries(payload))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    if mapping_path is not None:
        mapping_path.parent.mkdir(parents=True, exist_ok=True)
        payload["sanitized_file"] = str(output_path)
        write_mapping_data(mapping_path, payload)


def restore_docx(input_path: Path, output_path: Path, mapping_path: Path) -> None:
    if input_path.suffix.lower() != ".docx":
        raise ValueError("当前仅支持 .docx 文件。")
    payload = read_mapping(mapping_path)
    items = mapping_entries(payload, only_enabled=False)
    if not items:
        raise ValueError("映射文件中未找到有效 entries。")
    doc = Document(str(input_path))
    apply_replacements(doc, items, reverse=True)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log(f"还原完成: {output_path}")


def read_mapping(mapping_path: Path) -> dict[str, Any]:
    data = json.loads(mapping_path.read_text(encoding="utf-8"))
    if "entries" in data:
        entries = data["entries"]
    else:
        replacements = data.get("replacements", {})
        categories = data.get("categories", {})
        entries = [
            {
                "placeholder": placeholder,
                "original": str(original),
                "category": str(categories.get(placeholder, "AUTO")),
                "enabled": True,
                "source": "auto",
            }
            for placeholder, original in replacements.items()
        ]
    payload = {
        "version": 2,
        "source_file": data.get("source_file", ""),
        "sanitized_file": data.get("sanitized_file", ""),
        "entries": normalize_entries(entries),
    }
    refresh_payload_metadata(payload)
    return payload


def write_mapping_data(mapping_path: Path, payload: dict[str, Any]) -> None:
    refresh_payload_metadata(payload)
    mapping_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def build_mapping_payload(
    doc: Document,
    input_path: Path,
    output_path: Path,
    mapping_path: Path,
    custom_terms: list[str],
    use_llm_assist: bool,
    model: str,
    ollama_url: str,
    timeout_sec: int,
    retries: int,
    existing_entries_override: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    existing_entries: list[dict[str, Any]] = []
    if existing_entries_override is not None:
        existing_entries = normalize_entries(existing_entries_override)
        log(f"已加载当前内存映射: {len(existing_entries)} 条")
    elif mapping_path.exists():
        try:
            existing_entries = read_mapping(mapping_path).get("entries", [])
            log(f"已加载现有映射: {len(existing_entries)} 条")
        except Exception as exc:
            log(f"读取现有映射失败，将重新生成: {exc}", level="WARN")

    candidates = collect_candidates(doc, custom_terms)
    if use_llm_assist:
        texts = collect_doc_texts(doc)
        try:
            llm_candidates = collect_llm_candidates(texts, model=model, ollama_url=ollama_url, timeout_sec=timeout_sec, retries=retries)
            if llm_candidates:
                log(f"AI 辅助新增候选: {len(llm_candidates)} 条")
                candidates.extend(llm_candidates)
        except Exception as exc:
            log(f"AI 辅助识别失败，已回退规则模式: {exc}", level="WARN")
    merged = merge_entries(existing_entries, candidates)
    payload = {
        "version": 2,
        "source_file": str(input_path),
        "sanitized_file": str(output_path),
        "entries": merged,
    }
    refresh_payload_metadata(payload)
    log(f"识别到敏感项: {len(payload['entries'])} 条，启用 {len(payload['replacements'])} 条")
    return payload


def mapping_entries(payload: dict[str, Any], only_enabled: bool = True) -> list[ReplacementItem]:
    entries = normalize_entries(payload.get("entries", []))
    items: list[ReplacementItem] = []
    for entry in entries:
        if only_enabled and not entry["enabled"]:
            continue
        items.append(
            ReplacementItem(
                placeholder=entry["placeholder"],
                original=entry["original"],
                category=entry["category"],
                enabled=bool(entry.get("enabled", True)),
                source=str(entry.get("source", "auto")),
            )
        )
    return items


def normalize_entries(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen_pairs: set[tuple[str, str]] = set()
    for raw in entries:
        placeholder = normalize_text(str(raw.get("placeholder", "")))
        original = normalize_text(str(raw.get("original", "")))
        category = normalize_text(str(raw.get("category", "AUTO"))) or "AUTO"
        enabled = bool(raw.get("enabled", True))
        source = normalize_text(str(raw.get("source", "auto"))) or "auto"
        if not placeholder or not original:
            continue
        key = (placeholder, original)
        if key in seen_pairs:
            continue
        seen_pairs.add(key)
        out.append(
            {
                "placeholder": placeholder,
                "original": original,
                "category": category.upper(),
                "enabled": enabled,
                "source": source,
            }
        )
    return out


def refresh_payload_metadata(payload: dict[str, Any]) -> None:
    entries = normalize_entries(payload.get("entries", []))
    payload["entries"] = entries
    enabled_entries = [entry for entry in entries if entry["enabled"]]
    counts: dict[str, int] = {}
    for entry in enabled_entries:
        counts[entry["category"]] = counts.get(entry["category"], 0) + 1
    payload["counts"] = counts
    payload["replacements"] = {entry["placeholder"]: entry["original"] for entry in enabled_entries}
    payload["categories"] = {entry["placeholder"]: entry["category"] for entry in enabled_entries}


def merge_entries(existing_entries: list[dict[str, Any]], candidates: list[tuple[str, str, str]]) -> list[dict[str, Any]]:
    merged = normalize_entries(existing_entries)
    existing_by_original = {normalize_text(entry["original"]): entry for entry in merged}
    used_placeholders = {entry["placeholder"] for entry in merged}
    counters = category_counters(merged)

    for category, value, source in candidates:
        normalized = normalize_text(value)
        if normalized in existing_by_original:
            continue
        placeholder = next_placeholder(category, counters, used_placeholders)
        entry = {
            "placeholder": placeholder,
            "original": normalized,
            "category": category,
            "enabled": True,
            "source": source,
        }
        merged.append(entry)
        existing_by_original[normalized] = entry
    return merged


def collect_candidates(doc: Document, custom_terms: list[str]) -> list[tuple[str, str, str]]:
    texts = collect_doc_texts(doc)
    candidates: list[tuple[str, str, str]] = []

    for term in sorted({normalize_text(x) for x in custom_terms if normalize_text(x)}, key=len, reverse=True):
        candidates.append(("CUSTOM", term, "custom_terms"))

    for text in texts:
        for category, pattern in PLACEHOLDER_PATTERNS:
            for match in re.finditer(pattern, text):
                value = clean_candidate_value(normalize_text(match.group(0)), category)
                if is_valid_candidate(value, category):
                    candidates.append((category, value, "auto"))
        for category, pattern in CONTEXT_RULES:
            for match in re.finditer(pattern, text):
                value = clean_candidate_value(normalize_text(match.group(0)), category)
                if is_valid_candidate(value, category):
                    candidates.append((category, value, "auto"))
        for category, pattern in PERSON_CONTEXT_PATTERNS:
            for match in re.finditer(pattern, text):
                value = clean_candidate_value(normalize_text(match.group(1)), category)
                if is_valid_candidate(value, category):
                    candidates.append((category, value, "auto"))
        candidates.extend(extract_contextual_candidates(text))

    deduped: dict[str, tuple[str, str]] = {}
    for category, value, source in sorted(candidates, key=lambda item: len(item[1]), reverse=True):
        deduped.setdefault(normalize_text(value), (category, source))
    return [(category, original, source) for original, (category, source) in deduped.items()]


def collect_llm_candidates(
    texts: list[str],
    model: str,
    ollama_url: str,
    timeout_sec: int,
    retries: int,
) -> list[tuple[str, str, str]]:
    chunks = chunk_texts_for_llm(texts, max_chars=1800, max_items=16)
    candidates: list[tuple[str, str, str]] = []
    for idx, chunk in enumerate(chunks, start=1):
        log(f"AI 辅助识别分段 {idx}/{len(chunks)}")
        payload = call_ollama_candidate_json(
            ollama_url=ollama_url,
            model=model,
            prompt=build_llm_candidate_prompt(chunk),
            timeout_sec=timeout_sec,
            retries=retries,
        )
        for row in payload.get("candidates", []):
            if not isinstance(row, dict):
                continue
            category = normalize_text(str(row.get("category", "")).upper()) or "MANUAL"
            original = clean_candidate_value(normalize_text(str(row.get("text", ""))), category)
            if is_valid_candidate(original, category):
                candidates.append((category, original, "llm"))
    deduped: dict[str, tuple[str, str]] = {}
    for category, value, source in sorted(candidates, key=lambda item: len(item[1]), reverse=True):
        deduped.setdefault(normalize_text(value), (category, source))
    return [(category, original, source) for original, (category, source) in deduped.items()]


def chunk_texts_for_llm(texts: list[str], max_chars: int = 1800, max_items: int = 16) -> list[list[str]]:
    chunks: list[list[str]] = []
    current: list[str] = []
    current_len = 0
    for text in texts:
        line = normalize_text(text)
        if not line:
            continue
        line_len = len(line) + 1
        if current and (current_len + line_len > max_chars or len(current) >= max_items):
            chunks.append(current)
            current = []
            current_len = 0
        current.append(line)
        current_len += line_len
    if current:
        chunks.append(current)
    return chunks[:12]


def build_llm_candidate_prompt(chunk: list[str]) -> str:
    numbered = "\n".join(f"{idx + 1}. {line}" for idx, line in enumerate(chunk))
    return f"""请从下面文本中识别需要脱敏的敏感实体，并输出严格 JSON。

规则：
1. 只抽取文本中真实出现的原文，不改写，不杜撰。
2. 优先识别：公司主体、英文公司名、律所、人名、项目名、案号、合同编号、客户、供应商、代码、标题。
3. 只返回短实体，不要返回整句描述。
4. 如果不确定，不要返回。
5. category 只能是：COMPANY、PERSON、PROJECT、CASE、CODE、CUSTOMER、SUPPLIER、TITLE。
6. 像 Baker Botts、TransPerfect Legal、贝克博茨律所 这类律所或英文机构，也按 COMPANY 返回。

输出格式：
{{
  "candidates": [
    {{"text":"立讯技术股份有限公司","category":"COMPANY"}},
    {{"text":"张三","category":"PERSON"}}
  ]
}}

待识别文本：
{numbered}
"""


def extract_json_object(text: str) -> str:
    payload = text.strip()
    if payload.startswith("{") and payload.endswith("}"):
        return payload
    match = re.search(r"\{[\s\S]*\}", payload)
    if match:
        return match.group(0)
    raise ValueError("LLM 返回中未找到 JSON。")


def call_ollama_candidate_json(
    ollama_url: str,
    model: str,
    prompt: str,
    timeout_sec: int,
    retries: int,
) -> dict[str, Any]:
    system_prompt = "你是企业文档脱敏助手。只能输出严格 JSON，只抽取明确敏感实体，禁止返回整句。"
    endpoints = [
        ("chat", f"{ollama_url.rstrip('/')}/api/chat"),
        ("generate", f"{ollama_url.rstrip('/')}/api/generate"),
    ]
    last_error: Exception | None = None
    for attempt in range(1, retries + 1):
        for kind, url in endpoints:
            try:
                log(f"调用 Ollama 脱敏辅助 ({kind}) 第 {attempt} 次: {url}")
                if kind == "chat":
                    body = json.dumps(
                        {
                            "model": model,
                            "format": "json",
                            "stream": False,
                            "messages": [
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": prompt},
                            ],
                        }
                    ).encode("utf-8")
                else:
                    body = json.dumps(
                        {
                            "model": model,
                            "format": "json",
                            "stream": False,
                            "prompt": f"{system_prompt}\n\n{prompt}",
                        }
                    ).encode("utf-8")
                req = urllib.request.Request(url=url, data=body, headers={"Content-Type": "application/json"}, method="POST")
                with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
                    payload = json.loads(resp.read().decode("utf-8"))
                text = payload.get("message", {}).get("content", "") if kind == "chat" else payload.get("response", "")
                return json.loads(extract_json_object(text))
            except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, ValueError) as exc:
                last_error = exc
                log(f"Ollama 脱敏辅助失败: {exc}", level="WARN")
                continue
    if last_error is None:
        raise RuntimeError("Ollama 脱敏辅助失败。")
    raise RuntimeError(f"Ollama 脱敏辅助失败: {last_error}") from last_error


def extract_contextual_candidates(text: str) -> list[tuple[str, str, str]]:
    text = normalize_text(text)
    out: list[tuple[str, str, str]] = []
    for chunk in re.split(r"[，,。；;：:\s]+", text):
        chunk = normalize_text(chunk)
        if 4 <= len(chunk) <= 12 and re.fullmatch(r"[\u4e00-\u9fffA-Za-z0-9]{2,10}(?:集团|公司|科技|技术|电子|实业|贸易|国际)", chunk):
            chunk = clean_candidate_value(chunk, "COMPANY")
            if is_valid_candidate(chunk, "COMPANY"):
                out.append(("COMPANY", chunk, "auto"))
    if any(key in text for key in ["保密", "机密", "敏感", "涉美", "出口管制", "ECCN"]):
        for chunk in re.findall(r"[A-Z]{2,}(?:[-_][A-Z0-9]+)+", text):
            if is_valid_candidate(chunk, "CODE"):
                out.append(("CODE", chunk, "auto"))
    return out


def clean_candidate_value(value: str, category: str) -> str:
    value = normalize_text(value)
    if category == "PERSON" and len(value) >= 3 and value[-1] in {"非", "因", "与", "已", "将", "被", "系", "向", "于"}:
        value = value[:-1]
    if category == "CASE" and re.fullmatch(r"\d{4}", value):
        return ""
    return value


def collect_doc_texts(doc: Document) -> list[str]:
    texts: list[str] = []
    for paragraph in iter_doc_paragraphs(doc):
        text = normalize_text(paragraph.text)
        if text:
            texts.append(text)
    return texts


def iter_doc_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from iter_table_paragraphs(table)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from iter_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for inner in cell.tables:
                yield from iter_table_paragraphs(inner)


def apply_replacements(doc: Document, items: list[ReplacementItem], reverse: bool = False) -> None:
    ordered = sorted(items, key=lambda item: len(item.placeholder if reverse else item.original), reverse=True)
    for paragraph in iter_doc_paragraphs(doc):
        replace_in_paragraph(paragraph, ordered, reverse=reverse)


def replace_in_paragraph(paragraph, items: list[ReplacementItem], reverse: bool = False) -> None:
    source = paragraph.text or ""
    updated = source
    for item in items:
        old = item.placeholder if reverse else item.original
        new = item.original if reverse else item.placeholder
        if old in updated:
            updated = updated.replace(old, new)
    if updated == source:
        return
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = updated
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(updated)


def category_counters(entries: list[dict[str, Any]]) -> dict[str, int]:
    counters: dict[str, int] = {}
    for entry in entries:
        category = entry["category"]
        counters[category] = max(counters.get(category, 0), parse_placeholder_index(entry["placeholder"]))
    return counters


def parse_placeholder_index(placeholder: str) -> int:
    match = re.search(r"_(\d{3,})__", placeholder)
    return int(match.group(1)) if match else 0


def next_placeholder(category: str, counters: dict[str, int], used_placeholders: set[str]) -> str:
    category = category.upper()
    while True:
        counters[category] = counters.get(category, 0) + 1
        placeholder = f"__{category}_{counters[category]:03d}__"
        if placeholder not in used_placeholders:
            used_placeholders.add(placeholder)
            return placeholder


def make_manual_entry(original: str, placeholder: str | None = None, category: str = "MANUAL") -> dict[str, Any]:
    original = normalize_text(original)
    placeholder = normalize_text(placeholder or "")
    if not original:
        raise ValueError("敏感词不能为空。")
    if not placeholder:
        placeholder = f"__{category.upper()}_MANUAL__"
    return {
        "placeholder": placeholder,
        "original": original,
        "category": category.upper(),
        "enabled": True,
        "source": "manual",
    }


def is_valid_candidate(value: str, category: str) -> bool:
    value = normalize_text(value)
    if len(value) < 2:
        return False
    if any(char in value for char in "\n\t"):
        return False
    if sum(value.count(mark) for mark in "，,。；;！？!?") > 0:
        return False
    if category == "COMPANY":
        if value in GENERIC_COMPANY_VALUES:
            return False
        if len(value) < 4 or len(value) > 24:
            return False
        if any(token in value for token in SENTENCE_LIKE_TOKENS):
            return False
        if any(token in value for token in COMPANY_STOPWORDS):
            return False
        if re.match(r"^(?:自\d{4}|其中|协助|新增|需要|并约|申请人|被申请人|妥善|补充|提交|办理|联系|梳理|平台)", value):
            return False
        if value.endswith("公司") and len(value) <= 4:
            return False
    if category == "TITLE" and len(value.strip("《》")) < 2:
        return False
    if category == "ACCOUNT" and value.startswith("2025") and len(value) <= 12:
        return False
    if category == "PERSON":
        if len(value) > 4:
            return False
        if any(token in value for token in ["公司", "集团", "部门", "项目"]):
            return False
    if category in {"PARTY", "PROJECT", "SUPPLIER", "CUSTOMER"} and len(value) > 40:
        return False
    return True
